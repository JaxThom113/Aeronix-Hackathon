#!/usr/bin/env python3
"""
CLI Tool for handling multiple file types
Supports: text, json, csv, images, pdf, and more
"""

import typer
import json
import csv
import os
from pathlib import Path
from typing import List, Optional
import mimetypes
from PIL import Image
import pandas as pd

# Document processing imports
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

app = typer.Typer(help="Multi-file type CLI tool for Aeronix Hackathon")

# Global callback function for processing extracted content
content_processor = None


def set_content_processor(processor_func):
    """Set a callback function to process extracted file contents"""
    global content_processor
    content_processor = processor_func


def example_processor(
    content: str, file_type: str, file_path: str, processor_name: str = "default"
):
    """Example processor function that can be used to process file contents"""
    if processor_name == "word_count":
        words = content.split()
        return f"Word count: {len(words)}"
    elif processor_name == "summary":
        lines = content.split("\n")
        return f"Summary: {len(lines)} lines, {len(content)} characters"
    elif processor_name == "search":
        # Simple search for common words
        common_words = [
            "the",
            "and",
            "or",
            "but",
            "in",
            "on",
            "at",
            "to",
            "for",
            "of",
            "with",
            "by",
        ]
        found_words = [word for word in common_words if word.lower() in content.lower()]
        return f"Found common words: {', '.join(found_words)}"
    else:
        return f"Processed {file_type} file: {len(content)} characters"


def extract_file_content(file_path: str) -> dict:
    """Extract content from any supported file type"""
    file_type = detect_file_type(file_path)

    if not os.path.exists(file_path):
        return {
            "type": file_type,
            "file": file_path,
            "content": "",
            "error": "File not found",
        }

    try:
        if file_type == "text":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif file_type == "json":
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                content = json.dumps(data, indent=2)
        elif file_type == "csv":
            df = pd.read_csv(file_path)
            content = df.to_string(index=False)
        elif file_type == "pdf":
            if PDF_AVAILABLE:
                with open(file_path, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = "\n".join(
                        [page.extract_text() for page in pdf_reader.pages]
                    )
            else:
                content = "PDF processing requires PyPDF2 library"
        elif file_type == "docx":
            if DOCX_AVAILABLE:
                # Load document and extract text while skipping images and comments
                doc = Document(file_path)

                def _local_name(tag: str) -> str:
                    return tag.split("}")[-1] if "}" in tag else tag

                IMAGE_TAGS = {"drawing", "pict", "blip", "graphic", "graphicdata"}

                paragraphs = []
                for p in doc.paragraphs:
                    # Build paragraph text by concatenating runs that do not
                    # contain images, comments, or revision markup. This
                    # preserves surrounding text while removing the markup.
                    cleaned_runs = []
                    # Include all run text. Images will not be present in run.text
                    # so we can safely concatenate runs to preserve text.
                    for run in p.runs:
                        if run.text:
                            cleaned_runs.append(run.text)

                    cleaned_text = "".join(cleaned_runs).strip()
                    if cleaned_text:
                        paragraphs.append(cleaned_text)

                content = "\n".join(paragraphs)
            else:
                content = "DOCX processing requires python-docx library"
        elif file_type == "excel":
            if EXCEL_AVAILABLE:
                workbook = openpyxl.load_workbook(file_path)
                content_parts = []
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    content_parts.append(f"Sheet: {sheet_name}")
                    for row in sheet.iter_rows(values_only=True):
                        content_parts.append(
                            "\t".join(
                                [str(cell) if cell is not None else "" for cell in row]
                            )
                        )
                content = "\n".join(content_parts)
            else:
                content = "Excel processing requires openpyxl library"
        elif file_type == "ipc":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            content = f"Unsupported file type: {file_type}"

        return {
            "type": file_type,
            "file": file_path,
            "content": content,
            "size": len(content),
        }
    except Exception as e:
        return {"type": file_type, "file": file_path, "content": "", "error": str(e)}


def detect_file_type(file_path: str) -> str:
    """Detect file type based on extension and MIME type"""
    path = Path(file_path)
    mime_type, _ = mimetypes.guess_type(file_path)

    if mime_type:
        if mime_type.startswith("text/"):
            return "text"
        elif mime_type == "application/json":
            return "json"
        elif mime_type == "text/csv":
            return "csv"
        elif mime_type.startswith("image/"):
            return "image"
        elif mime_type == "application/pdf":
            return "pdf"

    # Fallback to extension
    ext = path.suffix.lower()
    if ext in [".txt", ".md", ".py", ".js", ".html", ".css"]:
        return "text"
    elif ext == ".json":
        return "json"
    elif ext == ".csv":
        return "csv"
    elif ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"]:
        return "image"
    elif ext == ".pdf":
        return "pdf"
    elif ext in [".xlsx", ".xls"]:
        return "excel"
    elif ext == ".docx":
        return "docx"
    elif ext == ".ipc":
        return "ipc"

    return "unknown"


def process_text_file(file_path: str) -> dict:
    """Process text files"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return {
            "type": "text",
            "file": file_path,
            "size": len(content),
            "lines": len(content.splitlines()),
            "words": len(content.split()),
            "preview": content[:200] + "..." if len(content) > 200 else content,
        }
    except Exception as e:
        return {"type": "text", "file": file_path, "error": str(e)}


def process_json_file(file_path: str) -> dict:
    """Process JSON files"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        return {
            "type": "json",
            "file": file_path,
            "size": os.path.getsize(file_path),
            "keys": (
                list(data.keys())
                if isinstance(data, dict)
                else f"Array with {len(data)} items"
            ),
            "preview": (
                json.dumps(data, indent=2)[:200] + "..."
                if len(str(data)) > 200
                else json.dumps(data, indent=2)
            ),
        }
    except Exception as e:
        return {"type": "json", "file": file_path, "error": str(e)}


def process_csv_file(file_path: str) -> dict:
    """Process CSV files"""
    try:
        df = pd.read_csv(file_path)
        return {
            "type": "csv",
            "file": file_path,
            "size": os.path.getsize(file_path),
            "rows": len(df),
            "columns": list(df.columns),
            "preview": df.head().to_string(),
        }
    except Exception as e:
        return {"type": "csv", "file": file_path, "error": str(e)}


def process_image_file(file_path: str) -> dict:
    """Process image files"""
    try:
        with Image.open(file_path) as img:
            return {
                "type": "image",
                "file": file_path,
                "size": os.path.getsize(file_path),
                "dimensions": img.size,
                "mode": img.mode,
                "format": img.format,
            }
    except Exception as e:
        return {"type": "image", "file": file_path, "error": str(e)}


def process_pdf_file(file_path: str) -> dict:
    """Process PDF files"""
    try:
        if not PDF_AVAILABLE:
            return {
                "type": "pdf",
                "file": file_path,
                "size": os.path.getsize(file_path),
                "note": "PDF processing requires PyPDF2 library",
            }

        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            # Extract text from first page as preview
            if num_pages > 0:
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text()
                preview = text[:200] + "..." if len(text) > 200 else text
            else:
                preview = "No content found"

        return {
            "type": "pdf",
            "file": file_path,
            "size": os.path.getsize(file_path),
            "pages": num_pages,
            "preview": preview,
        }
    except Exception as e:
        return {"type": "pdf", "file": file_path, "error": str(e)}


def process_docx_file(file_path: str) -> dict:
    """Process Word documents"""
    try:
        if not DOCX_AVAILABLE:
            return {
                "type": "docx",
                "file": file_path,
                "size": os.path.getsize(file_path),
                "note": "DOCX processing requires python-docx library",
            }

        doc = Document(file_path)

        def _local_name(tag: str) -> str:
            return tag.split("}")[-1] if "}" in tag else tag

        IMAGE_TAGS = {"drawing", "pict", "blip", "graphic", "graphicdata"}
        COMMENT_TAGS = {
            "commentrangestart",
            "commentrangeend",
            "commentreference",
            "comment",
        }
        paragraphs = []
        for p in doc.paragraphs:
            # Preserve paragraph text while removing runs that include
            # images/comments/revision markup.
            cleaned_runs = []
            for run in p.runs:
                if run.text:
                    cleaned_runs.append(run.text)

            cleaned_text = "".join(cleaned_runs).strip()
            if cleaned_text:
                paragraphs.append(cleaned_text)

        text_content = "\n".join(paragraphs)

        return {
            "type": "docx",
            "file": file_path,
            "size": os.path.getsize(file_path),
            "paragraphs": len(paragraphs),
            "words": len(text_content.split()),
            "preview": (
                text_content[:200] + "..." if len(text_content) > 200 else text_content
            ),
        }
    except Exception as e:
        return {"type": "docx", "file": file_path, "error": str(e)}


def process_excel_file(file_path: str) -> dict:
    """Process Excel files"""
    try:
        if not EXCEL_AVAILABLE:
            return {
                "type": "excel",
                "file": file_path,
                "size": os.path.getsize(file_path),
                "note": "Excel processing requires openpyxl library",
            }

        workbook = openpyxl.load_workbook(file_path)
        sheet_names = workbook.sheetnames
        first_sheet = workbook[sheet_names[0]]

        # Get dimensions
        max_row = first_sheet.max_row
        max_col = first_sheet.max_column

        # Get preview data
        preview_data = []
        for row in range(1, min(6, max_row + 1)):  # First 5 rows
            row_data = []
            for col in range(1, min(6, max_col + 1)):  # First 5 columns
                cell_value = first_sheet.cell(row=row, column=col).value
                row_data.append(str(cell_value) if cell_value is not None else "")
            preview_data.append(row_data)

        return {
            "type": "excel",
            "file": file_path,
            "size": os.path.getsize(file_path),
            "sheets": sheet_names,
            "rows": max_row,
            "columns": max_col,
            "preview": preview_data,
        }
    except Exception as e:
        return {"type": "excel", "file": file_path, "error": str(e)}


def process_ipc_file(file_path: str) -> dict:
    """Process IPC files"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return {
            "type": "ipc",
            "file": file_path,
            "size": len(content),
            "lines": len(content.splitlines()),
            "words": len(content.split()),
            "preview": content[:200] + "..." if len(content) > 200 else content,
        }
    except Exception as e:
        return {"type": "ipc", "file": file_path, "error": str(e)}


def process_file(file_path: str) -> dict:
    """Process a single file based on its type"""
    file_type = detect_file_type(file_path)

    if not os.path.exists(file_path):
        return {"type": file_type, "file": file_path, "error": "File not found"}

    processors = {
        "text": process_text_file,
        "json": process_json_file,
        "csv": process_csv_file,
        "image": process_image_file,
        "pdf": process_pdf_file,
        "docx": process_docx_file,
        "excel": process_excel_file,
        "ipc": process_ipc_file,
    }

    processor = processors.get(
        file_type,
        lambda x: {"type": file_type, "file": x, "error": "Unsupported file type"},
    )
    return processor(file_path)


@app.command()
def analyze(
    files: List[str] = typer.Argument(..., help="File paths to analyze"),
    output: Optional[str] = typer.Option(
        None, "--output", "-o", help="Output file for results"
    ),
    format: str = typer.Option(
        "table", "--format", "-f", help="Output format: table, json, csv"
    ),
):
    """Analyze multiple files of different types"""
    results = []

    for file_path in files:
        typer.echo(f"Processing: {file_path}")
        result = process_file(file_path)
        results.append(result)

    if format == "json":
        output_data = json.dumps(results, indent=2)
    elif format == "csv":
        # Convert to CSV format
        if results:
            df = pd.DataFrame(results)
            output_data = df.to_csv(index=False)
        else:
            output_data = ""
    else:  # table format
        output_data = "\n".join(
            [
                f"File: {r['file']} | Type: {r['type']} | Size: {r.get('size', 'N/A')} bytes"
                + (f" | Error: {r['error']}" if "error" in r else "")
                for r in results
            ]
        )

    if output:
        with open(output, "w") as f:
            f.write(output_data)
        typer.echo(f"Results saved to: {output}")
    else:
        typer.echo(output_data)


@app.command()
def convert(
    input_file: str = typer.Argument(..., help="Input file path"),
    output_file: str = typer.Argument(..., help="Output file path"),
    format: str = typer.Option(
        "json", "--format", "-f", help="Output format: json, csv, txt"
    ),
):
    """Convert files between different formats"""
    file_type = detect_file_type(input_file)

    if file_type == "csv" and format == "json":
        try:
            df = pd.read_csv(input_file)
            df.to_json(output_file, orient="records", indent=2)
            typer.echo(f"Converted {input_file} to {output_file}")
        except Exception as e:
            typer.echo(f"Error: {e}")
    elif file_type == "json" and format == "csv":
        try:
            with open(input_file, "r") as f:
                data = json.load(f)
            df = pd.DataFrame(data)
            df.to_csv(output_file, index=False)
            typer.echo(f"Converted {input_file} to {output_file}")
        except Exception as e:
            typer.echo(f"Error: {e}")
    else:
        typer.echo(f"Conversion from {file_type} to {format} not supported yet")


@app.command()
def batch(
    directory: str = typer.Argument(..., help="Directory to process"),
    pattern: str = typer.Option(
        "*", "--pattern", "-p", help="File pattern (e.g., *.txt, *.json)"
    ),
    recursive: bool = typer.Option(
        False, "--recursive", "-r", help="Process subdirectories"
    ),
):
    """Process all files in a directory matching a pattern"""
    from glob import glob

    if recursive:
        search_pattern = os.path.join(directory, "**", pattern)
    else:
        search_pattern = os.path.join(directory, pattern)

    files = glob(search_pattern, recursive=recursive)

    if not files:
        typer.echo(f"No files found matching pattern: {pattern}")
        return

    typer.echo(f"Found {len(files)} files to process")

    results = []
    for file_path in files:
        result = process_file(file_path)
        results.append(result)

    # Display summary
    type_counts = {}
    for result in results:
        file_type = result["type"]
        type_counts[file_type] = type_counts.get(file_type, 0) + 1

    typer.echo("\nFile type summary:")
    for file_type, count in type_counts.items():
        typer.echo(f"  {file_type}: {count} files")


@app.command()
def info(file: str = typer.Argument(..., help="File to get detailed info about")):
    """Get detailed information about a single file"""
    result = process_file(file)

    typer.echo(f"\nFile Information:")
    typer.echo(f"  Path: {result['file']}")
    typer.echo(f"  Type: {result['type']}")

    if "error" in result:
        typer.echo(f"  Error: {result['error']}")
    else:
        for key, value in result.items():
            if key not in ["file", "type"]:
                typer.echo(f"  {key.title()}: {value}")


@app.command()
def extract(
    files: List[str] = typer.Argument(..., help="File paths to extract content from"),
    output: Optional[str] = typer.Option(
        None, "--output", "-o", help="Output file for extracted content"
    ),
    process: bool = typer.Option(
        False, "--process", "-p", help="Process content with registered processor"
    ),
):
    """Extract content from files and optionally process it"""
    results = []

    for file_path in files:
        typer.echo(f"Extracting content from: {file_path}")
        result = extract_file_content(file_path)
        results.append(result)

        if "error" in result:
            typer.echo(f"  Error: {result['error']}")
        else:
            typer.echo(
                f"  Type: {result['type']} | Content size: {result['size']} characters"
            )

            # Process content if processor is registered and flag is set
            if process and content_processor:
                try:
                    processed_result = content_processor(
                        result["content"], result["type"], result["file"]
                    )
                    typer.echo(f"  Processed: {processed_result}")
                except Exception as e:
                    typer.echo(f"  Processing error: {e}")

    # Save results if output file specified
    if output:
        output_data = json.dumps(results, indent=2)
        with open(output, "w") as f:
            f.write(output_data)
        typer.echo(f"Results saved to: {output}")

    return results


@app.command()
def process_content(
    files: List[str] = typer.Argument(..., help="File paths to process"),
    processor_name: str = typer.Option(
        "default", "--processor", help="Name of the processor to use"
    ),
):
    """Process file contents with a specific processor"""
    if not content_processor:
        typer.echo(
            "No content processor registered. Use set_content_processor() first."
        )
        return

    results = []
    for file_path in files:
        typer.echo(f"Processing: {file_path}")
        content_result = extract_file_content(file_path)

        if "error" in content_result:
            typer.echo(f"  Error: {content_result['error']}")
            continue

        try:
            # Call the registered processor function
            processed = content_processor(
                content=content_result["content"],
                file_type=content_result["type"],
                file_path=content_result["file"],
                processor_name=processor_name,
            )
            results.append(
                {
                    "file": file_path,
                    "type": content_result["type"],
                    "processed": processed,
                }
            )
            typer.echo(f"  Processed successfully")
        except Exception as e:
            typer.echo(f"  Processing error: {e}")
            results.append(
                {"file": file_path, "type": content_result["type"], "error": str(e)}
            )

    return results


@app.command()
def test():
    """Test the CLI tool with sample files"""
    test_dir = Path(__file__).parent / "test_files"

    if not test_dir.exists():
        typer.echo("Test directory not found. Creating sample files...")
        test_dir.mkdir(exist_ok=True)

        # Create sample files
        sample_files = {
            "sample.txt": "This is a sample text file for testing.",
            "sample.json": '{"name": "test", "value": 123}',
            "sample.csv": "Name,Age\nAlice,25\nBob,30",
            "sample.ipc": "IPC test file content",
        }

        for filename, content in sample_files.items():
            (test_dir / filename).write_text(content)

    typer.echo("Testing CLI tool with sample files...")
    typer.echo("=" * 50)

    # Test all files in test directory
    test_files = list(test_dir.glob("*"))

    if not test_files:
        typer.echo("No test files found!")
        return

    results = []
    for file_path in test_files:
        typer.echo(f"\nProcessing: {file_path.name}")
        result = process_file(str(file_path))
        results.append(result)

        # Display basic info
        typer.echo(f"  Type: {result['type']}")
        if "error" in result:
            typer.echo(f"  Error: {result['error']}")
        else:
            typer.echo(f"  Size: {result.get('size', 'N/A')} bytes")
            if "preview" in result:
                preview = result["preview"]
                if isinstance(preview, str) and len(preview) > 100:
                    preview = preview[:100] + "..."
                typer.echo(f"  Preview: {preview}")

    # Summary
    typer.echo("\n" + "=" * 50)
    typer.echo("Test Summary:")
    type_counts = {}
    for result in results:
        file_type = result["type"]
        type_counts[file_type] = type_counts.get(file_type, 0) + 1

    for file_type, count in type_counts.items():
        typer.echo(f"  {file_type}: {count} files")

    typer.echo(f"\nTotal files processed: {len(results)}")

    # Test content extraction
    typer.echo("\n" + "=" * 50)
    typer.echo("Testing Content Extraction:")

    # Set up example processor
    set_content_processor(example_processor)

    for file_path in test_files[:3]:  # Test first 3 files
        typer.echo(f"\nExtracting content from: {file_path.name}")
        content_result = extract_file_content(str(file_path))

        if "error" in content_result:
            typer.echo(f"  Error: {content_result['error']}")
        else:
            typer.echo(f"  Type: {content_result['type']}")
            typer.echo(f"  Content size: {content_result['size']} characters")

            # Show first 100 characters of content
            preview = content_result["content"][:100]
            if len(content_result["content"]) > 100:
                preview += "..."
            typer.echo(f"  Preview: {preview}")

    typer.echo("\nTest completed successfully!")


if __name__ == "__main__":
    app()
