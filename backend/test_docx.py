#!/usr/bin/env python3
"""
Test DOCX generation function
"""

import sys
import os
from pathlib import Path

# Add backend to path
backend_dir = Path(__file__).parent
sys.path.insert(0, str(backend_dir))


def test_docx_creation():
    """Test DOCX creation function"""
    try:
        from docx import Document
        from datetime import datetime
        import tempfile

        print("Testing DOCX creation...")

        # Create a simple document
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")

        # Add metadata
        doc.add_heading("Document Information", level=1)
        doc.add_paragraph(
            f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        )
        doc.add_paragraph("Content length: 50 characters")
        doc.add_paragraph("Word count: 10 words")

        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, "test_output.docx")
        print(f"Saving to: {output_path}")

        doc.save(output_path)
        print(f"DOCX created successfully: {output_path}")
        print(f"File exists: {os.path.exists(output_path)}")
        print(f"File size: {os.path.getsize(output_path)} bytes")

        return output_path

    except Exception as e:
        print(f"Error creating DOCX: {e}")
        import traceback

        traceback.print_exc()
        return None


if __name__ == "__main__":
    test_docx_creation()
