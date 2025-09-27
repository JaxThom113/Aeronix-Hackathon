#!/usr/bin/env python3
"""
Test script for the Flask API
"""

import requests
import json
import os
import shutil
from pathlib import Path

API_BASE_URL = "http://localhost:5000"
TEST_OUTPUT_DIR = "test_outputs"


def clear_test_outputs():
    """Clear the test output directory"""
    if os.path.exists(TEST_OUTPUT_DIR):
        shutil.rmtree(TEST_OUTPUT_DIR)
    os.makedirs(TEST_OUTPUT_DIR, exist_ok=True)
    print(f"Cleared test output directory: {TEST_OUTPUT_DIR}")


def test_health():
    """Test health endpoint"""
    print("Testing health endpoint...")
    response = requests.get(f"{API_BASE_URL}/health")
    print(f"Status: {response.status_code}")
    print(f"Response: {response.json()}")
    print()


def test_upload_file():
    """Test file upload endpoint"""
    print("Testing file upload...")

    # Create a test file
    test_file_path = "backend/test_files/sample.txt"
    if not os.path.exists(test_file_path):
        print(f"Test file {test_file_path} not found!")
        return

    with open(test_file_path, "rb") as f:
        files = {"file": f}
        response = requests.post(f"{API_BASE_URL}/upload", files=files)

    print(f"Status: {response.status_code}")
    if response.status_code == 200:
        result = response.json()
        print(f"Success: {result['success']}")
        print(f"Filename: {result['filename']}")
        print(f"File type: {result['result']['type']}")
        print(f"Content size: {result['result']['size']} characters")
    else:
        print(f"Error: {response.json()}")
    print()


def test_process_files():
    """Test file processing endpoint"""
    print("Testing file processing...")

    # Test with local files
    test_files = [
        "backend/test_files/sample.txt",
        "backend/test_files/sample.json",
        "backend/test_files/sample.csv",
    ]

    # Filter existing files
    existing_files = [f for f in test_files if os.path.exists(f)]

    if not existing_files:
        print("No test files found!")
        return

    data = {"files": existing_files, "processor": "word_count"}

    response = requests.post(f"{API_BASE_URL}/process", json=data)

    print(f"Status: {response.status_code}")
    if response.status_code == 200:
        result = response.json()
        print(f"Success: {result['success']}")
        print(f"Processor used: {result['processor_used']}")
        print(f"Files processed: {len(result['results'])}")

        for file_result in result["results"]:
            print(f"  {file_result['file']}: {file_result['type']}")
            if "processed" in file_result:
                print(f"    Processed: {file_result['processed']}")
    else:
        print(f"Error: {response.json()}")
    print()


def test_gemini_text():
    """Test Gemini text processing"""
    print("Testing Gemini text processing...")

    test_text = "This is a sample document for testing the Gemini API integration. It contains multiple sentences and should be processed by the mock Gemini service."

    data = {"text": test_text}
    response = requests.post(f"{API_BASE_URL}/gemini-text", json=data)

    print(f"Status: {response.status_code}")
    if response.status_code == 200:
        result = response.json()
        print(f"Success: {result['success']}")
        print(f"Original text length: {result['response_length']}")
        print(f"Gemini response preview: {result['gemini_response'][:200]}...")
    else:
        print(f"Error: {response.json()}")
    print()


def test_gemini_docx():
    """Test Gemini DOCX generation"""
    test_text = "This is a comprehensive document that will be processed by Gemini and converted to a DOCX file. The content includes various sections and should demonstrate the full workflow."

    data = {"text": test_text}
    response = requests.post(f"{API_BASE_URL}/gemini", json=data)

    print(f"Status: {response.status_code}")
    if response.status_code == 200:
        # Save the DOCX file
        output_path = os.path.join(TEST_OUTPUT_DIR, "test_gemini_output.docx")
        with open(output_path, "wb") as f:
            f.write(response.content)
        print(f"DOCX file saved as: {output_path}")
        print(f"File size: {os.path.getsize(output_path)} bytes")
    else:
        print(f"Error: {response.json()}")
    print()


def test_upload_and_gemini():
    """Test complete workflow: upload file -> Gemini -> DOCX"""
    print("Testing complete workflow...")

    # Use a test file
    test_file_path = "backend/test_files/sample.txt"
    if not os.path.exists(test_file_path):
        print(f"Test file {test_file_path} not found!")
        return

    with open(test_file_path, "rb") as f:
        files = {"file": f}
        response = requests.post(f"{API_BASE_URL}/upload-and-gemini", files=files)

    print(f"Status: {response.status_code}")
    if response.status_code == 200:
        # Save the DOCX file
        output_path = os.path.join(TEST_OUTPUT_DIR, "test_upload_gemini_output.docx")
        with open(output_path, "wb") as f:
            f.write(response.content)
        print(f"DOCX file saved as: {output_path}")
        print(f"File size: {os.path.getsize(output_path)} bytes")
    else:
        print(f"Error: {response.json()}")
    print()


def test_upload_multiple():
    """Test uploading multiple files and receiving file info list"""
    print("Testing upload multiple files...")

    # Prepare test files
    file_paths = [
        "backend/test_files/sample.txt",
        "backend/test_files/sample.json",
    ]

    files = []
    for p in file_paths:
        if not os.path.exists(p):
            print(f"Test file {p} not found, skipping test.")
            return
        files.append(("files", open(p, "rb")))

    try:
        response = requests.post(f"{API_BASE_URL}/upload-multiple", files=files)
    finally:
        # Close opened files
        for _, fh in files:
            try:
                fh.close()
            except Exception:
                pass

    print(f"Status: {response.status_code}")
    assert response.status_code == 200
    data = response.json()
    print(f"Response: {data}")
    assert data.get("success") is True

    print()


def test_merge_docx():
    """Test merging two DOCX files and verifying AI-generated marker"""
    print("Testing merge DOCX...")

    doc1 = "backend/test_files/Clemson_HW_Spec_V4_092325.docx"
    doc2 = "backend/test_files/Clemson_SW_Spec 2.1.docx"

    for p in (doc1, doc2):
        if not os.path.exists(p):
            print(f"Test doc {p} not found, skipping test.")
            return

    files = [
        ("files", open(doc1, "rb")),
        ("files", open(doc2, "rb")),
    ]

    response = None
    try:
        try:
            response = requests.post(f"{API_BASE_URL}/merge-docx", files=files)
        except requests.exceptions.ConnectionError:
            # Server not reachable via HTTP; fall back to Flask test client
            print("Server not reachable via HTTP; falling back to Flask test client")
            import sys

            # Ensure backend is importable
            sys.path.insert(0, os.path.join(os.getcwd(), "backend"))
            import api as local_api

            client = local_api.app.test_client()

            # Close request-opened file objects and re-open for test client
            for _, fh in files:
                try:
                    fh.close()
                except Exception:
                    pass

            opened = []
            # Use a MultiDict so EnvironBuilder can accept multiple files with the same key
            from werkzeug.datastructures import MultiDict

            data = MultiDict()
            for p in (doc1, doc2):
                f = open(p, "rb")
                opened.append(f)
                # value is a (fileobj, filename) tuple
                data.add("files", (f, os.path.basename(p)))

            try:
                response = client.post(
                    "/merge-docx", data=data, content_type="multipart/form-data"
                )
            finally:
                for f in opened:
                    try:
                        f.close()
                    except Exception:
                        pass
    finally:
        # Ensure any files opened for the requests call are closed
        for _, fh in files:
            try:
                fh.close()
            except Exception:
                pass

    output_path = os.path.join(TEST_OUTPUT_DIR, "merged.docx")
    with open(output_path, "wb") as f:
        f.write(response.content)

    print(f"Merged doc saved to: {output_path}")

    # Verify merged content contains the AI marker heading
    try:
        from docx import Document

        merged = Document(output_path)
        first_para = merged.paragraphs[0].text if merged.paragraphs else ""
        print(f"First paragraph: {first_para}")
        assert "AI-Generated" in first_para
    except ImportError:
        print("python-docx not installed for test verification")

    print()


def main():
    """Run all tests"""
    print("=== Flask API Test Suite ===")
    print()

    # Clear test outputs at the beginning
    clear_test_outputs()
    print()

    try:
        test_health()
        test_upload_file()
        test_process_files()
        test_gemini_text()
        test_gemini_docx()
        test_upload_and_gemini()
        test_upload_multiple()
        test_merge_docx()

        print("All tests completed!")
        print(f"Test output files saved in: {TEST_OUTPUT_DIR}/")

    except requests.exceptions.ConnectionError:
        print(
            "Error: Could not connect to API. Make sure the Flask server is running on localhost:5000"
        )
        print("Start the server with: python backend/api.py")
    except Exception as e:
        print(f"Error during testing: {e}")


if __name__ == "__main__":
    main()
