from datetime import datetime
import io
import os
import tempfile
from collections import defaultdict
import subprocess
import tempfile


from flask import Flask, request, jsonify, send_file
from docx import Document
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain_google_genai import ChatGoogleGenerativeAI

from cli import extract_file_content
from netlist_parser import build_pipeline, parse_d356, parse_ipc

app = Flask(__name__)
qa_chain = None

# ---- your helpers (assumed to exist) ----
# parse_netlist, parse_d356, extract_file_content must already be defined


def flatten_netlist(parsed):
    snippets = []

    # d356-like dict with 'components' and 'nets'
    if "components" in parsed and "nets" in parsed:
        if parsed.get("metadata"):
            meta_chunk = "Metadata:\n" + "\n".join(
                [f"{k}: {v}" for k, v in parsed["metadata"].items()]
            )
            snippets.append(meta_chunk)

        for comp_name, comp_data in parsed.get("components", {}).items():
            pin_lines = []
            for pin in comp_data.get("pins", []):
                pin_lines.append(
                    f"Pin connected to net '{pin['net_name']}' (ID: {pin['net_id']}) at ({pin['x']}, {pin['y']}), rotation {pin['rotation']}, symbol {pin['symbol_id']}"
                )
            comp_chunk = f"Component: {comp_name}\n" + "\n".join(pin_lines)
            snippets.append(comp_chunk)

        for net_id, net_data in parsed.get("nets", {}).items():
            connections = net_data.get("connections", [])
            net_chunk = f"Net: {net_id} ({net_data.get('name', '')})\nConnected components: {', '.join(connections)}"
            snippets.append(net_chunk)

    # ipc-style list of dicts
    elif isinstance(parsed, list):
        net_groups = defaultdict(list)
        for entry in parsed:
            net_groups[entry["net"]].append(entry)

        for net, entries in net_groups.items():
            lines = []
            for e in entries:
                lines.append(
                    f"Component {e['component']}, Pin {e['pin']}, Pos ({e['x']},{e['y']}), Rot {e['rotation']}, Side {e['side']}"
                )
            net_chunk = f"Net: {net}\n" + "\n".join(lines)
            snippets.append(net_chunk)

    return snippets


def create_docx_bytes(text, original_content=None) -> bytes:
    """
    Create a DOCX as bytes (no temp files).
    Accepts any 'text' (will be str()'d), normalizes newlines,
    and builds a simple, Word-friendly document.
    """
    # Coerce to strings and normalize newlines
    text = "" if text is None else str(text)
    original_content = None if original_content is None else str(original_content)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if original_content is not None:
        original_content = original_content.replace("\r\n", "\n").replace("\r", "\n")

    doc = Document()

    # Title
    doc.add_heading("AI-Generated Document", 0)

    # Original content (optional)
    if original_content:
        doc.add_heading("Original Input Content", level=1)
        doc.add_paragraph("The following is the original content that was processed:")
        doc.add_paragraph("─" * 50)
        for line in original_content.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph("─" * 50)
        doc.add_paragraph("")

    # Processed content
    doc.add_heading("AI Analysis & Processing Results", level=1)
    # Split on blank lines into paragraphs
    for para in [p.strip() for p in text.split("\n\n") if p.strip()]:
        doc.add_paragraph(para)

    # Metadata
    doc.add_heading("Document Information", level=1)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f"Content length: {len(text)} characters")
    doc.add_paragraph(f"Word count: {len(text.split())} words")
    if original_content:
        doc.add_paragraph(
            f"Original content length: {len(original_content)} characters"
        )
        doc.add_paragraph(f"Original word count: {len(original_content.split())} words")

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# Put this near the top if you want a fixed project root (repo root):
PROJECT_ROOT = os.path.abspath(
    os.getcwd()
)  # assume you're running from Aeronix-Hackathon root


def _resolve_local_paths(paths):
    """Resolve client-sent paths safely to local absolute paths under PROJECT_ROOT (or accept absolute paths)."""
    resolved = []
    for p in paths:
        # Accept absolute paths as-is; otherwise, treat as relative to PROJECT_ROOT
        abs_path = os.path.abspath(
            p if os.path.isabs(p) else os.path.join(PROJECT_ROOT, p)
        )
        # Prevent path traversal outside the project root for relative inputs
        if not os.path.isabs(p) and not abs_path.startswith(PROJECT_ROOT):
            raise ValueError(f"Unsafe path outside project root: {p}")
        if not os.path.exists(abs_path) or not os.path.isfile(abs_path):
            raise FileNotFoundError(f"File not found: {p}")
        resolved.append(abs_path)
    return resolved


@app.route("/upload", methods=["POST"])
def upload_files():
    global qa_chain

    data = request.get_json(silent=True) or {}
    requested_paths = data.get("files", [])
    if not isinstance(requested_paths, list) or not requested_paths:
        return (
            jsonify(
                {
                    "error": 'Send JSON like {"files": ["backend/test_files/foo.ipc", ...]}'
                }
            ),
            400,
        )

    try:
        file_paths = _resolve_local_paths(requested_paths)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    # Build entries [(path, ext, name)] for downstream logic
    entries = []
    for path in file_paths:
        name = os.path.basename(path)
        ext = os.path.splitext(name)[1].lower()
        entries.append((path, ext, name))

    snippets: list[str] = []
    processed: list[dict] = []

    # --- Parse/group content (unchanged logic, but using 'entries' instead of temp files) ---
    for file_path, ext, orig_name in entries:
        try:
            if ext == ".d356":
                parsed = parse_d356(file_path)
                snippets.extend(flatten_netlist(parsed))
                processed.append(
                    {
                        "file": orig_name,
                        "type": "d356",
                        "status": "ok",
                        "snippets_added": "netlist",
                    }
                )

            elif ext == ".ipc":
                try:
                    parsed = parse_ipc(file_path)
                except NameError:
                    parsed = extract_file_content(file_path)
                if isinstance(parsed, (dict, list)):
                    snippets.extend(flatten_netlist(parsed))
                    processed.append(
                        {
                            "file": orig_name,
                            "type": "ipc(parsed)",
                            "status": "ok",
                            "snippets_added": "netlist",
                        }
                    )
                else:
                    snippets.append(str(parsed))
                    processed.append(
                        {
                            "file": orig_name,
                            "type": "ipc(text)",
                            "status": "ok",
                            "snippets_added": "text",
                        }
                    )

            elif ext == ".docx":
                d = Document(file_path)
                doc_text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
                if doc_text.strip():
                    snippets.append(doc_text)
                processed.append(
                    {
                        "file": orig_name,
                        "type": "docx",
                        "status": "ok",
                        "snippets_added": "text",
                    }
                )

            else:
                # generic text fallback
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as rf:
                        txt = rf.read()
                        if txt.strip():
                            snippets.append(txt)
                    processed.append(
                        {
                            "file": orig_name,
                            "type": ext or "unknown",
                            "status": "ok",
                            "snippets_added": "text",
                        }
                    )
                except Exception as e:
                    processed.append(
                        {
                            "file": orig_name,
                            "type": ext or "unknown",
                            "status": f"skipped ({e})",
                            "snippets_added": 0,
                        }
                    )

        except Exception as e:
            processed.append(
                {
                    "file": orig_name,
                    "type": ext or "unknown",
                    "status": f"error ({e})",
                    "snippets_added": 0,
                }
            )

    if not snippets:
        return (
            jsonify(
                {
                    "status": "no_content",
                    "message": "No parsable text extracted.",
                    "files_processed": processed,
                    "snippets_count": 0,
                }
            ),
            400,
        )

    # --- Embeddings + Vectorstore ---
    embedder = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vectorstore = Chroma.from_texts(
        texts=snippets, embedding=embedder, persist_directory="./chroma_db"
    )
    vectorstore.persist()

    # --- LLM + RetrievalQA (GLOBAL) ---
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash")
    retriever = vectorstore.as_retriever(
        search_type="similarity", search_kwargs={"k": 143}
    )
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm, chain_type="stuff", retriever=retriever
    )

    return jsonify(
        {"status": "ok", "files_processed": processed, "snippets_count": len(snippets)}
    )


@app.route("/download")
def get_file():
    global qa_chain
    question = """You are a useful tool used by engineers to generate bring-up test plans. Use the provided files, which may include a netlist, BOM,
                  hardware requirement document, and software requirement document. Additionally, reference the
                  provided example bring-up test document for guidance. Make the formatting clean and professional. Try to
                  format it so it works perfectly for word documents and looks nice. Use headings, bullet points,
                  and numbered lists where appropriate. Include a title page with generation date and title of what is being accomplished.
                  """
    text = qa_chain.run(question)

    if not text:
        return jsonify({"error": "Missing required field: 'text'"}), 400

    filename = "TestProcedure.docx"

    doc_bytes = create_docx_bytes(text)
    # doc_bytes = markdown_to_docx_bytes(text)
    return send_file(
        io.BytesIO(doc_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,  # Flask 2.x
    )


def markdown_to_docx_bytes(md_text: str) -> bytes:
    """Convert Markdown (or plain text) to DOCX via Pandoc and return bytes."""
    md_text = "" if md_text is None else str(md_text)
    # Pandoc handles plain text too; no special formatting required.
    with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as md_file:
        md_file.write(md_text.encode("utf-8"))
        md_path = md_file.name

    out_fd, out_path = tempfile.mkstemp(suffix=".docx")
    os.close(out_fd)

    try:
        subprocess.run(["pandoc", md_path, "-o", out_path], check=True)
        with open(out_path, "rb") as f:
            return f.read()
    finally:
        for p in (md_path, out_path):
            try:
                os.remove(p)
            except Exception:
                pass


if __name__ == "__main__":
    app.run(debug=True)
